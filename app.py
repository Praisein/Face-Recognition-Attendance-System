import cv2
from flask_cors import CORS
from mobile_routes import register_mobile_routes
from flask import (
    Flask, render_template, request, jsonify, Response,
    redirect, url_for, session, flash, send_file)
import threading
import multiprocessing
import json
import time
import os
from datetime import datetime
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from io import BytesIO
from openpyxl import Workbook
from docx import Document
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from functools import wraps
# import attendance system utilities
import json
import os
from datetime import datetime
from attendance_system import main as attendance_main, atomic_write_json, load_student_data, save_student_data
from curriculum_toggle import get_state as get_curriculum_state, toggle_curriculum
# Import IP access control
from ip_access_control import (
    enable_mobile_access, disable_mobile_access, 
    get_access_status, is_mobile_access_enabled
)

try:
    from curriculum_toggle import get_state as get_curriculum_state, toggle as toggle_curriculum
except Exception:
    # If import fails, provide fallbacks
    def get_curriculum_state():
        return {'current_side': 'odd', 'odd_true': 0, 'even_true': 0, 'semesters_total': 0}
    def toggle_curriculum():
        return 'odd'
from ip_access_control import (
    check_mobile_access, 
    is_localhost, 
    is_network_device,
    get_lan_ip,
    get_access_status
)
from flask import render_template_string, abort

# -----------------------------
# SECTION: Flask app initialization & basic config
# (create Flask app, secret key, process/thread placeholders)
# -----------------------------

app = Flask(__name__)
app.secret_key = os.urandom(24)  # change in production
process_thread = None


# Add this after creating the Flask app instance (after line 34)
# app = Flask(__name__)
# -----------------------------
# SECTION: CORS and security configuration
# (Cross-Origin Resource Sharing, Content Security Policy headers)
# -----------------------------

CORS(app, resources={
    r"/*": {
        "origins": "*",
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"]
    }
})

# -----------------------------
# SECTION: Network & access control constants
# (server LAN IP, mobile/public route lists)
# -----------------------------

SERVER_LAN_IP = get_lan_ip()

# Routes that network devices can access ONLY when mobile access is enabled
MOBILE_ONLY_ROUTES = [
    '/',
    '/mobile_attendance',  # The main mobile attendance page
    '/mobile_recognize'    # Backend API to process face recognition
]

# Routes that are always accessible to everyone (even without login)
PUBLIC_ROUTES = [
    '/login',
    '/logout',
    '/verify_teacher'
]

# -----------------------------
# SECTION: IP access enforcement middleware
# (blocks or allows requests based on device IP and mobile access)
# -----------------------------

@app.before_request
def enforce_ip_access():
    request_path = request.path
    client_ip = request.remote_addr  # <-- Fix: define client_ip from request
    if any(request_path.startswith(route) for route in PUBLIC_ROUTES):
        return None
    if is_localhost(client_ip):
        return None
    if any(request_path.startswith(route) for route in MOBILE_ONLY_ROUTES):
        allowed, reason = check_mobile_access(client_ip)
        
        if not allowed:
            # Log the denied access attempt
            print(f"[ACCESS DENIED] {client_ip} → {request_path} (Reason: {reason})")
            
            # Return friendly error page
            return render_template_string('''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>🔒 Mobile Access Disabled</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            display: flex;
            align-items: center;
            justify-content: center;
            min-height: 100vh;
            padding: 20px;
        }
        .container {
            background: white;
            padding: 40px;
            border-radius: 16px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            max-width: 500px;
            text-align: center;
            animation: slideUp 0.5s ease;
        }
        @keyframes slideUp {
            from {
                opacity: 0;
                transform: translateY(30px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        .icon {
            font-size: 80px;
            margin-bottom: 20px;
            animation: pulse 2s infinite;
        }
        @keyframes pulse {
            0%, 100% { transform: scale(1); }
            50% { transform: scale(1.1); }
        }
        h1 {
            color: #dc2626;
            font-size: 28px;
            margin-bottom: 16px;
        }
        p {
            color: #374151;
            line-height: 1.8;
            margin-bottom: 16px;
            font-size: 16px;
        }
        .info-box {
            background: #f3f4f6;
            padding: 16px;
            border-radius: 8px;
            margin: 24px 0;
            border-left: 4px solid #667eea;
        }
        .info-box strong {
            color: #374151;
            display: block;
            margin-bottom: 4px;
        }
        .info-box code {
            color: #6b7280;
            font-family: 'Courier New', monospace;
            font-size: 14px;
            word-break: break-all;
        }
        .instructions {
            background: #fffbeb;
            border: 2px solid #fbbf24;
            padding: 20px;
            border-radius: 8px;
            margin: 24px 0;
            text-align: left;
        }
        .instructions h3 {
            color: #d97706;
            margin-bottom: 12px;
            font-size: 18px;
        }
        .instructions ol {
            color: #374151;
            padding-left: 20px;
            line-height: 2;
        }
        .btn {
            display: inline-block;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 14px 32px;
            border-radius: 8px;
            text-decoration: none;
            font-weight: 600;
            margin-top: 20px;
            transition: transform 0.2s;
            box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4);
        }
        .btn:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(102, 126, 234, 0.6);
        }
        .countdown {
            margin-top: 20px;
            font-size: 14px;
            color: #6b7280;
        }
        .countdown strong {
            color: #667eea;
            font-size: 18px;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="icon">🔒</div>
        <h1>Mobile Access Disabled</h1>
        <p><strong>This page is currently not accessible from your device.</strong></p>
        <p>Mobile attendance access must be enabled by your teacher.</p>
        
        <div class="info-box">
            <strong>Your Device IP:</strong>
            <code>{{ client_ip }}</code>
            <br><br>
            <strong>Server IP:</strong>
            <code>{{ server_ip }}</code>
        </div>
        
        <div class="instructions">
            <h3>📋 Teacher Instructions:</h3>
            <ol>
                <li>Open the app on your laptop</li>
                <li>Go to <strong>Mark Attendance</strong> page</li>
                <li>Click <strong>"Enable Mobile Access"</strong> button</li>
                <li>Share the URL with students</li>
            </ol>
        </div>
        
        <p style="font-size: 14px; color: #6b7280; margin-top: 24px;">
            Once enabled, access will be available for <strong>5 minutes</strong>.
        </p>
        
        <a href="javascript:location.reload()" class="btn">🔄 Retry Now</a>
        
        <div class="countdown">
            Auto-retrying in <strong id="countdown">5</strong> seconds...
        </div>
    </div>
    
    <script>
        // Auto-retry countdown
        let seconds = 5;
        const countdownEl = document.getElementById('countdown');
        
        const timer = setInterval(() => {
            seconds--;
            countdownEl.textContent = seconds;
            
            if (seconds <= 0) {
                clearInterval(timer);
                location.reload();
            }
        }, 1000);
    </script>
</body>
</html>
            ''', client_ip=client_ip, server_ip=SERVER_LAN_IP), 403
        
        # If allowed, proceed to the route
        return None
    
    # ========================================================================
    # RULE 4: Network devices trying to access other pages (admin pages)
    # ========================================================================
    if is_network_device(client_ip):
        # Log the blocked attempt
        print(f"[ACCESS BLOCKED] {client_ip} → {request_path} (Not a mobile route)")
        
        # Return strict denial page
        return render_template_string('''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>⛔ Access Denied</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: #1f2937;
            display: flex;
            align-items: center;
            justify-content: center;
            min-height: 100vh;
            padding: 20px;
            color: white;
        }
        .container {
            background: #374151;
            padding: 40px;
            border-radius: 16px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.5);
            max-width: 500px;
            text-align: center;
            border: 3px solid #dc2626;
            animation: shake 0.5s;
        }
        @keyframes shake {
            0%, 100% { transform: translateX(0); }
            25% { transform: translateX(-10px); }
            75% { transform: translateX(10px); }
        }
        .icon {
            font-size: 100px;
            margin-bottom: 20px;
            filter: drop-shadow(0 0 20px rgba(220, 38, 38, 0.5));
        }
        h1 {
            color: #fca5a5;
            font-size: 32px;
            margin-bottom: 16px;
        }
        p {
            color: #d1d5db;
            line-height: 1.8;
            margin-bottom: 16px;
            font-size: 16px;
        }
        .info-box {
            background: #1f2937;
            padding: 20px;
            border-radius: 8px;
            margin: 24px 0;
            border: 1px solid #4b5563;
        }
        .info-box strong {
            color: #fca5a5;
            display: block;
            margin-bottom: 8px;
        }
        .info-box code {
            color: #9ca3af;
            font-family: 'Courier New', monospace;
            font-size: 14px;
            word-break: break-all;
        }
        .warning {
            background: #7f1d1d;
            border: 2px solid #dc2626;
            padding: 20px;
            border-radius: 8px;
            margin-top: 24px;
        }
        .warning h3 {
            color: #fca5a5;
            margin-bottom: 12px;
        }
        .warning p {
            font-size: 14px;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="icon">⛔</div>
        <h1>Access Denied</h1>
        <p><strong>This page cannot be accessed from network devices.</strong></p>
        <p>The system can only be fully accessed from the host computer (localhost).</p>
        
        <div class="info-box">
            <strong>Your IP:</strong>
            <code>{{ client_ip }}</code>
            <br><br>
            <strong>Requested Page:</strong>
            <code>{{ path }}</code>
            <br><br>
            <strong>Access Level:</strong>
            <code>Network Device (Restricted)</code>
        </div>
        
        <div class="warning">
            <h3>🔐 Security Notice</h3>
            <p>
                Administrative pages, teacher controls, and student data 
                can only be accessed from the host device for security reasons.
            </p>
            <p style="margin-top: 12px;">
                Network devices can only access mobile attendance features 
                when explicitly enabled by the teacher.
            </p>
        </div>
    </div>
</body>
</html>
        ''', client_ip=client_ip, path=request_path), 403
    
    # ========================================================================
    # FALLBACK: Allow request (shouldn't reach here normally)
    # ========================================================================
    return None


# ============================================================================
# OPTIONAL: Server Info Route (Localhost Only)
# ============================================================================

# -----------------------------
# SECTION: Authentication helper: login_required decorator
# (used to protect routes requiring teacher login)
# -----------------------------

def login_required(f):
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session and 'teacher_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/server_info')
@login_required  # Ensure user is logged in
def server_info():
    """
    Display server connection information.
    Only accessible from localhost.
    """
    # Extra security: block if not localhost
    if not is_localhost(request.remote_addr):
        abort(403, "This page is only accessible from localhost")
    
    access_status = get_access_status()
    
    return render_template_string('''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>🖥️ Server Information</title>
    <style>
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: #f8fafc;
            padding: 40px 20px;
        }
        .container {
            max-width: 900px;
            margin: 0 auto;
        }
        .card {
            background: white;
            padding: 30px;
            border-radius: 12px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
            margin-bottom: 24px;
        }
        h1 {
            color: #374151;
            margin-bottom: 24px;
        }
        h2 {
            color: #667eea;
            font-size: 20px;
            margin-bottom: 16px;
            border-bottom: 2px solid #e5e7eb;
            padding-bottom: 8px;
        }
        .info-grid {
            display: grid;
            gap: 12px;
        }
        .info-item {
            display: flex;
            justify-content: space-between;
            padding: 12px;
            background: #f9fafb;
            border-radius: 6px;
            border-left: 4px solid #667eea;
        }
        .label {
            font-weight: 600;
            color: #374151;
        }
        .value {
            font-family: 'Courier New', monospace;
            color: #6b7280;
            word-break: break-all;
            text-align: right;
        }
        .status-enabled {
            color: #059669;
            font-weight: 700;
        }
        .status-disabled {
            color: #dc2626;
            font-weight: 700;
        }
        .btn {
            display: inline-block;
            padding: 12px 24px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            text-decoration: none;
            border-radius: 8px;
            margin-right: 12px;
            margin-top: 16px;
            font-weight: 600;
            transition: transform 0.2s;
        }
        .btn:hover {
            transform: translateY(-2px);
        }
        .btn-secondary {
            background: #6b7280;
        }
        ol {
            line-height: 2;
            color: #374151;
            padding-left: 20px;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="card">
            <h1>🖥️ Server Information</h1>
            
            <h2>📡 Network Details</h2>
            <div class="info-grid">
                <div class="info-item">
                    <span class="label">Localhost URL:</span>
                    <span class="value">http://127.0.0.1:5000</span>
                </div>
                <div class="info-item">
                    <span class="label">Network URL:</span>
                    <span class="value">http://{{ lan_ip }}:5000</span>
                </div>
                <div class="info-item">
                    <span class="label">Mobile Attendance URL:</span>
                    <span class="value">http://{{ lan_ip }}:5000/mobile_attendance</span>
                </div>
            </div>
            
            <h2>📱 Mobile Access Status</h2>
            <div class="info-grid">
                <div class="info-item">
                    <span class="label">Status:</span>
                    <span class="value {% if status.enabled %}status-enabled{% else %}status-disabled{% endif %}">
                        {% if status.enabled %}✅ ENABLED{% else %}❌ DISABLED{% endif %}
                    </span>
                </div>
                {% if status.enabled %}
                <div class="info-item">
                    <span class="label">Time Remaining:</span>
                    <span class="value status-enabled">{{ status.remaining_formatted }}</span>
                </div>
                <div class="info-item">
                    <span class="label">Expires At:</span>
                    <span class="value">{{ status.expiry_time }}</span>
                </div>
                {% endif %}
            </div>
            
            <h2>📋 How to Use Mobile Access</h2>
            <ol>
                <li>Go to <strong>Mark Attendance</strong> page on this laptop</li>
                <li>Click the <strong>"Enable Mobile Access"</strong> button</li>
                <li>Copy the mobile URL and share it with students</li>
                <li>Students access the URL from their phones (same Wi-Fi)</li>
                <li>They take a selfie to mark attendance</li>
                <li>Access automatically expires after 5 minutes</li>
            </ol>
            
            <a href="/mark_attendance" class="btn">📸 Go to Mark Attendance</a>
            <a href="/dashboard" class="btn btn-secondary">🏠 Back to Dashboard</a>
        </div>
    </div>
</body>
</html>
    ''', lan_ip=SERVER_LAN_IP, status=access_status)


#clear session data
@app.after_request
def add_header(response):
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session and 'teacher_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function
#till here

# Add CSP configuration for Bootstrap and external resources
@app.after_request
def set_response_headers(response):
    response.headers['Content-Security-Policy'] = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' https://code.jquery.com https://cdn.jsdelivr.net "
        "https://stackpath.bootstrapcdn.com https://cdnjs.cloudflare.com; "
        "style-src 'self' 'unsafe-inline' https://stackpath.bootstrapcdn.com "
        "https://cdnjs.cloudflare.com https://fonts.googleapis.com; "
        "font-src 'self' https://fonts.gstatic.com https://cdnjs.cloudflare.com; "
        "img-src 'self' data: blob:; "
        "connect-src 'self';"
    )
    return response

CURRENT_TEACHER_JSON = "current_teacher.json"
VERIFICATION_JSON = "teacher_verification.json"  # ✅ Added verification file
os.makedirs("static/teacher_images", exist_ok=True)
os.makedirs("static/student_images", exist_ok=True)

# JSON storage filenames
ATTENDANCE_RECORDS_JSON = "attendance_records.json"
STUDENT_DATA_JSON = "student_data.json"
TEACHERS_JSON = "teachers.json"
# Ensure these are defined before init_database uses them
TEACHER_DATA_JSON = "teacher_data.json"
CURRICULUM_JSON = "curriculum.json"


# -----------------------------
# SECTION: Data persistence helpers (load/save/init)
# (functions that read/write JSON storage files)
# -----------------------------

def load_attendance_records():
    """Load attendance records from JSON file"""
    try:
        with open('attendance_records.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        return {'records': {}}
    except Exception as e:
        print(f"[ERROR] Failed to load attendance_records.json: {e}")
        return {'records': {}}

def save_attendance_records(data):
    """Save attendance records to JSON file"""
    try:
        atomic_write_json('attendance_records.json', data)
    except Exception as e:
        print(f"[ERROR] Failed to save attendance_records.json: {e}")

def init_database():
    """Initialize JSON files if they don't exist"""
    if not os.path.exists(ATTENDANCE_RECORDS_JSON):
        atomic_write_json(ATTENDANCE_RECORDS_JSON, {'records': {}})
    if not os.path.exists(STUDENT_DATA_JSON):
        atomic_write_json(STUDENT_DATA_JSON, {'students': {}})
    if not os.path.exists(TEACHER_DATA_JSON):
        atomic_write_json(TEACHER_DATA_JSON, {})
    if not os.path.exists(CURRICULUM_JSON):
        atomic_write_json(CURRICULUM_JSON, {})
    print("[INFO] JSON files initialized successfully")


def load_teachers():
    try:
        with open(TEACHER_DATA_JSON, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        return {}
    except Exception as e:
        print("[ERROR] Failed to load teachers.json:", e)
        return {}

def load_curriculum_data():
    # Load your curriculum JSON file
    # Adjust the path to match your file structure
    try:
        with open('curriculum.json', 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        app.logger.error('Curriculum JSON file not found')
        return {}
    except json.JSONDecodeError:
        app.logger.error('Invalid JSON in curriculum file')
        return {}

def load_verification_data():
    try:
        with open(VERIFICATION_JSON, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        print("[WARN] teacher_verification.json not found.")
        return {}
    except Exception as e:
        print("[ERROR] Failed to load teacher_verification.json:", e)
        return {}


# ----- New: load teacher_data.json (used by app1.py) -----
TEACHER_DATA_JSON = "teacher_data.json"

def load_teacher_data():
    try:
        with open(TEACHER_DATA_JSON, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        return {}
    except Exception as e:
        print("[ERROR] Failed to load teacher_data.json:", e)
        return {}


def find_teacher(teacher_id, password):
    try:
        td = load_teacher_data() or {}

        def _normalize(s):
            if not isinstance(s, str):
                return ''
            s2 = s.strip().lower()
            for ch in ['\u2013', '\u2014', '\u2010', '\u2011', '\u2012']:
                s2 = s2.replace(ch, '-')
            s2 = s2.replace('_', '-').replace(' ', '').replace('\u00A0', '')
            return s2

        for dept, teachers in td.items():
            if not isinstance(teachers, dict):
                continue

            # Direct or tolerant match
            found_key = None
            if teacher_id in teachers:
                found_key = teacher_id
            else:
                target_norm = _normalize(teacher_id)
                for k in teachers.keys():
                    if _normalize(k) == target_norm:
                        found_key = k
                        break

            if not found_key:
                continue

            entry = teachers.get(found_key, {})
            stored_password = entry.get('password', '')

            ok = False
            try:
                if stored_password and check_password_hash(stored_password, password):
                    ok = True
            except Exception:
                ok = False
            if not ok and stored_password == password:
                ok = True

            if ok:
                return {
                    'teacher_id': found_key,
                    'name': entry.get('name', ''),
                    'username': entry.get('username', found_key),
                    'subjects': entry.get('subjects', []),
                    'department': dept,
                    'password': stored_password
                }

        return None

    except Exception as e:
        print(f"[ERROR] find_teacher: {e}")
        return None


# Add this route to your app.py file (place it before the /signup route)

# -----------------------------
# SECTION: Authentication routes (verify/login/logout)
# (verify_teacher API, login/logout handlers)
# -----------------------------

@app.route('/verify_teacher', methods=['POST'])
def verify_teacher():
    """Verify teacher credentials against the JSON database"""
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        teacher_id = data.get('teacher_id', '').strip()

        if not username or not teacher_id:
            return jsonify({'verified': False, 'message': 'Missing credentials'})

        # Load verification data
        allowed_teachers = load_verification_data()

        # Check if teacher exists and credentials match
        if teacher_id in allowed_teachers:
            teacher_info = allowed_teachers[teacher_id]

            if teacher_info.get('username') == username:
                # Teacher verified - prepare response
                photo_url = None
                photo_path = os.path.join("static", "teacher_images", f"{teacher_id}.png")

                # Check if photo exists
                if os.path.exists(photo_path):
                    photo_url = url_for('static', filename=f"teacher_images/images/Teacher.jpg")

                return jsonify({
                    'verified': True,
                    'teacher_id': teacher_id,
                    'username': username,
                    'name': teacher_info.get('name', ''),
                    'photo_url': photo_url
                })

        return jsonify({'verified': False, 'message': 'Invalid credentials'})

    except Exception as e:
        print(f"[ERROR] Verification failed: {e}")
        return jsonify({'verified': False, 'message': 'Server error'}), 500

# ---------- AUTH (signup/login) ----------

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page supporting both teacher_data.json and teachers.json"""
    if request.method == 'POST':
        # Method 1: teacher_id/password (teacher_data.json)
        teacher_id = request.form.get('teacher_id', '').upper().strip()
        password = request.form.get('password', '').strip()

        if teacher_id and password:
            # Debug logging to help diagnose login issues
            try:
                print(f"[DEBUG] Login attempt: teacher_id={teacher_id}, password_provided={bool(password)}")
                td_sample = load_teacher_data()
                # print top-level dept keys and whether teacher_id exists under any
                dept_keys = list(td_sample.keys()) if isinstance(td_sample, dict) else []
                found_in = [d for d, tbl in td_sample.items() if isinstance(tbl, dict) and teacher_id in tbl]
                print(f"[DEBUG] teacher_data departments={dept_keys}; found_in={found_in}")
            except Exception as _:
                print("[DEBUG] Could not introspect teacher_data.json")

            teacher = find_teacher(teacher_id, password)
            print(f"[DEBUG] find_teacher returned: {bool(teacher)}")
            if teacher:
                session.clear()
                
                # Set all session variables
                session['teacher_id'] = teacher['teacher_id']
                session['teacher_name'] = teacher.get('name', '')
                session['teacher_subjects'] = teacher.get('subjects', [])
                session['department'] = teacher.get('department', '')
                session['username'] = teacher.get('username', teacher['teacher_id'])
                
                # Set default lecture
                if teacher.get('subjects'):
                    session['lecture'] = teacher['subjects'][0]
                    session['selected_subject'] = teacher['subjects'][0]

                # Handle teacher image
                teacher_image = f"{teacher_id}.png"
                teacher_image_path = os.path.join("static", "teacher_images", teacher_image)
                if os.path.exists(teacher_image_path):
                    session['teacher_image'] = teacher_image
                else:
                    session['teacher_image'] = None

                # Save current teacher
                atomic_write_json(CURRENT_TEACHER_JSON, {
                    "username": session['username'],
                    "name": session['teacher_name'],
                    "lecture": session.get('lecture', ''),
                    "image": session['teacher_image']
                })

                # Redirect based on number of subjects
                if len(teacher.get('subjects', [])) > 1:
                    return redirect(url_for('select_subject'))
                else:
                    return redirect(url_for('dashboard'))

        # Method 2: username/password (teachers.json) - fallback
        username = request.form.get('username', '').strip()
        password_fallback = request.form.get('password', '').strip()
        
        if username and password_fallback:
            teachers = load_teachers()
            if username in teachers and check_password_hash(teachers[username]['password'], password_fallback):
                session.clear()
                
                session['username'] = username
                session['teacher_name'] = teachers[username]['name']
                session['lecture'] = teachers[username]['lecture']
                session['selected_subject'] = teachers[username]['lecture']
                session['teacher_subjects'] = [teachers[username]['lecture']]
                session['teacher_image'] = teachers[username].get('photo')

                atomic_write_json(CURRENT_TEACHER_JSON, {
                    "username": username,
                    "name": teachers[username]['name'],
                    "lecture": teachers[username]['lecture'],
                    "image": session['teacher_image']
                })

                return redirect(url_for('dashboard'))

        # If we get here, credentials were invalid
        return render_template('login.html', error='Invalid Teacher ID or Password')

    # GET request
    return render_template('login.html')

def collect_teacher_data():
    """Collect currently logged-in teacher data from session and files"""
    if 'teacher_id' not in session and 'username' not in session:
        return None

    teacher_id = session.get('teacher_id')
    teacher_name = session.get('teacher_name', '')
    username = session.get('username', '')
    subjects = session.get('teacher_subjects', [])
    department = session.get('department', '')

    # Get teacher photo (session stores file name like 'IT-01.png')
    teacher_image = session.get('teacher_image')
    if teacher_image:
        photo_url = url_for('static', filename=f'teacher_images/{teacher_image}')
    else:
        # fallback placeholder image
        photo_url = "/images/Teacher.jpg"

    teacher_data = {
        'photo': photo_url,
        'name': teacher_name,
        'id': teacher_id or username,
        'username': username,
        'subjects': subjects,
        'department': department
    }

    return teacher_data

def collect_lecture_data():
    """Collect all lecture sessions (date, time, and name) for the selected subject."""
    from datetime import datetime

    subject = (session.get('selected_subject') or session.get('lecture') or '').strip()
    teacher_id = session.get('teacher_id', session.get('username', 'Unknown'))

    if not subject:
        return None

    try:
        with open('attendance_records.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {'teacher_id': teacher_id, 'subject': subject, 'total_lectures': 0, 'lectures': []}

    records = data.get('records', {})
    lectures = []

    subj_norm = subject.lower()

    for key, rec in records.items():
        if '_' not in key:
            continue

        date_part, lecture_name = key.split('_', 1)

        # Match the subject name (case-insensitive, exact match)
        if lecture_name.strip().lower() != subj_norm:
            continue

        lectures.append({
            'date': date_part,
            'time': rec.get('time', ''),
            'lecture_name': lecture_name
        })

    # Sort by date (latest first)
    try:
        lectures.sort(key=lambda x: datetime.strptime(x['date'], "%Y-%m-%d"), reverse=True)
    except Exception:
        pass

    return {
        'teacher_id': teacher_id,
        'subject': subject,
        'total_lectures': len(lectures),
        'lectures': lectures
    }

@app.route('/select-subject', methods=['GET', 'POST'])
@login_required
def select_subject():
    # Require login
    if 'teacher_id' not in session and 'username' not in session:
        return redirect(url_for('login'))
    
    # Load curriculum data from your JSON file
    curriculum = load_curriculum_data()  # You'll need to implement this function
    
    if request.method == 'POST':
        subject = request.form.get('subject')
        teacher_subjects = session.get('teacher_subjects') or []

        # Direct match
        if subject in teacher_subjects:
            session['selected_subject'] = subject
            session['lecture'] = subject
            # Persist current teacher selection so the attendance process sees it
            try:
                atomic_write_json(CURRENT_TEACHER_JSON, {
                    'username': session.get('username', ''),
                    'name': session.get('teacher_name', ''),
                    'lecture': subject,
                    'image': session.get('teacher_image')
                })
                print(f"[DEBUG] Wrote current_teacher.json with lecture='{subject}' for user={session.get('username')}")
            except Exception:
                app.logger.exception('Failed to write current_teacher.json')
            return redirect(url_for('dashboard'))

        # Case-insensitive match
        lower = subject.strip().lower() if subject else ''
        for s in teacher_subjects:
            if s and s.strip().lower() == lower:
                session['selected_subject'] = s
                session['lecture'] = s
                return redirect(url_for('dashboard'))

        # Fallback: reload teacher data by teacher_id and try matching there
        try:
            teacher_id = session.get('teacher_id')
            if teacher_id:
                td = load_teacher_data() or {}
                for dept, tbl in td.items():
                    if isinstance(tbl, dict) and teacher_id in tbl:
                        entry = tbl.get(teacher_id, {})
                        subjects = entry.get('subjects', [])
                        # update session cache
                        session['teacher_subjects'] = subjects
                        # direct or case-insensitive match
                        if subject in subjects:
                            session['selected_subject'] = subject
                            session['lecture'] = subject
                            try:
                                atomic_write_json(CURRENT_TEACHER_JSON, {
                                    'username': session.get('username', ''),
                                    'name': session.get('teacher_name', ''),
                                    'lecture': subject,
                                    'image': session.get('teacher_image')
                                })
                            except Exception:
                                app.logger.exception('Failed to write current_teacher.json')
                            return redirect(url_for('dashboard'))
                        for s in subjects:
                            if s and s.strip().lower() == lower:
                                session['selected_subject'] = s
                                session['lecture'] = s
                                try:
                                    atomic_write_json(CURRENT_TEACHER_JSON, {
                                        'username': session.get('username', ''),
                                        'name': session.get('teacher_name', ''),
                                        'lecture': s,
                                        'image': session.get('teacher_image')
                                    })
                                    print(f"[DEBUG] Wrote current_teacher.json with lecture='{s}' for user={session.get('username')}")
                                except Exception:
                                    app.logger.exception('Failed to write current_teacher.json')
                                return redirect(url_for('dashboard'))
                        break
        except Exception:
            app.logger.exception('Error while validating subject against teacher_data')

        # As a pragmatic fallback accept the submitted subject (trust the form) and update session
        if subject:
            subjects = session.get('teacher_subjects') or []
            if subject not in subjects:
                subjects.append(subject)
                session['teacher_subjects'] = subjects
            session['selected_subject'] = subject
            session['lecture'] = subject
            app.logger.warning('Accepted subject via fallback: %s', subject)
            try:
                atomic_write_json(CURRENT_TEACHER_JSON, {
                    'username': session.get('username', ''),
                    'name': session.get('teacher_name', ''),
                    'lecture': subject,
                    'image': session.get('teacher_image')
                })
                print(f"[DEBUG] Wrote current_teacher.json with lecture='{subject}' for user={session.get('username')}")
            except Exception:
                app.logger.exception('Failed to write current_teacher.json')
            return redirect(url_for('dashboard'))

        # If no subject provided, render page with error
        return render_template('select_subject.html',
                               teacher_name=session.get('teacher_name'),
                               curriculum=curriculum,
                               selected_year=request.args.get('year'),
                               selected_semester=request.args.get('semester'),
                               selected_type=request.args.get('subject_type'),
                               filtered_subjects=[],
                               error='Invalid subject selected')

    # GET request - handle filtering
    selected_year = request.args.get('year')
    selected_semester = request.args.get('semester')
    selected_type = request.args.get('subject_type')
    
    teacher_subjects = session.get('teacher_subjects', [])
    filtered_subjects = []

    # Apply filters if all required filters are selected
    if selected_year and selected_semester:
        semester_data = curriculum.get(selected_year, {}).get('Semesters', {}).get(selected_semester, {})
        
        if selected_type == 'Theory':
            subjects_from_curriculum = semester_data.get('Theory', [])
        elif selected_type == 'Practicals':
            subjects_from_curriculum = semester_data.get('Practicals', [])
        else:
            # Combine both Theory and Practicals if no type is selected
            subjects_from_curriculum = semester_data.get('Theory', []) + semester_data.get('Practicals', [])
        
        # Filter to only show subjects that the teacher has
        filtered_subjects = [subject for subject in subjects_from_curriculum if subject in teacher_subjects]
    
    return render_template('select_subject.html',
                           teacher_name=session.get('teacher_name'),
                           curriculum=curriculum,
                           selected_year=selected_year,
                           selected_semester=selected_semester,
                           selected_type=selected_type,
                           filtered_subjects=filtered_subjects)



# Add this new route for easy mobile access
@app.route('/mobile')
def mobile_redirect():
    """Redirect to mobile attendance page"""
    if 'username' not in session and 'teacher_id' not in session:
        return redirect(url_for('login'))
    return redirect(url_for('mobile_attendance')) or redirect(url_for('home'))

@app.route('/attendance')
def attendance():
    # Backwards-compatible attendance view
    if 'selected_subject' not in session and 'lecture' not in session:
        return redirect(url_for('select_subject'))

    # Initialize attendance records in session if not exists
    if 'attendance_records' not in session:
        session['attendance_records'] = []

    return render_template('attendance.html',
                           teacher_name=session.get('teacher_name'),
                           teacher_id=session.get('teacher_id', session.get('username', '')),
                           subject=session.get('selected_subject', session.get('lecture', '')),
                           records=session.get('attendance_records', []))

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


# ---------- DASHBOARD & MARK ATTENDANCE ----------
# -----------------------------
# SECTION: Dashboard & attendance UI routes
# (dashboard, mark_attendance, manage_students, exports)
# -----------------------------

@app.route('/dashboard')
def dashboard():
    if 'username' in session or 'teacher_id' in session:
        teacher_photo_url = None
        teacher_name = session.get('teacher_name', '')

        # Try to get teacher image from session
        if session.get('teacher_image'):
            teacher_photo_url = url_for('static', filename=f"teacher_images/{session['teacher_image']}")
        else:
            # Fallback to teachers.json lookup
            try:
                teachers = load_teachers()
                username = session.get('username')
                if username and username in teachers:
                    t = teachers[username]
                    if t and t.get('photo'):
                        teacher_photo_url = url_for('static', filename=f"teacher_images/{t['photo']}")
            except Exception:
                teacher_photo_url = None

        return render_template('home.html',
                               teacher_name=teacher_name,
                               teacher_photo=teacher_photo_url,
                               teacher_id=session.get('teacher_id'),
                               department=session.get('department'),
                               subject=session.get('selected_subject', session.get('lecture', '')))
    return redirect(url_for('login'))

@app.route('/teacher_profile')
@login_required
def teacher_profile():
    teacher_data = collect_teacher_data()
    if not teacher_data:
        return redirect(url_for('login'))

    return render_template('teacher_profile.html', teacher_data=teacher_data)

# Update the mark_attendance route to include mobile option
@app.route('/mark_attendance')
def mark_attendance():
    if 'username' in session or 'teacher_id' in session:
        # Detect mobile user agent
        user_agent = request.headers.get('User-Agent', '').lower()
        is_mobile = any(device in user_agent for device in ['mobile', 'android', 'iphone', 'ipad'])

        # Existing desktop code...
        teacher_photo_url = None
        if session.get('teacher_image'):
            teacher_photo_url = url_for('static', filename=f"teacher_images/{session['teacher_image']}")
        else:
            try:
                teachers = load_teachers()
                username = session.get('username')
                if username and username in teachers:
                    t = teachers[username]
                    if t and t.get('photo'):
                        teacher_photo_url = url_for('static', filename=f"teacher_images/{t['photo']}")
            except Exception:
                teacher_photo_url = None

        student_photo_url = None
        try:
            with open("current_student.json", "r", encoding='utf-8') as f:
                student_info = json.load(f)
                sid = student_info.get('student_id')
                if sid:
                    student_photo_url = url_for('static', filename=f"student_images/images/Student.jpg")
        except Exception:
            student_photo_url = None

        # Get mobile access status
        access_status = get_access_status()

        return render_template('mark_attendance.html',
                               teacher_name=session.get('teacher_name', ''),
                               username=session.get('username', ''),
                               teacher_photo=teacher_photo_url,
                               student_photo=student_photo_url,
                               teacher_id=session.get('teacher_id'),
                               department=session.get('department'),
                               subject=session.get('selected_subject', session.get('lecture', '')),
                               mobile_access_status=access_status)
    return redirect(url_for('login'))


# NEW ROUTES FOR MOBILE ACCESS CONTROL
@app.route('/api/mobile_access/toggle', methods=['POST'])
@login_required
def toggle_mobile_access():
    """Toggle mobile access on/off"""
    try:
        data = request.get_json() or {}
        action = data.get('action', 'toggle')
        
        if action == 'enable':
            duration = data.get('duration', 5)  # Default 5 minutes
            success, expiry_time = enable_mobile_access(duration)
            
            if success:
                return jsonify({
                    'success': True,
                    'enabled': True,
                    'message': f'Mobile access enabled for {duration} minutes',
                    'expiry_time': expiry_time.isoformat(),
                    'status': get_access_status()
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'Failed to enable mobile access'
                }), 500
                
        elif action == 'disable':
            success = disable_mobile_access()
            
            if success:
                return jsonify({
                    'success': True,
                    'enabled': False,
                    'message': 'Mobile access disabled',
                    'status': get_access_status()
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'Failed to disable mobile access'
                }), 500
        else:
            # Toggle based on current state
            if is_mobile_access_enabled():
                success = disable_mobile_access()
                return jsonify({
                    'success': success,
                    'enabled': False,
                    'message': 'Mobile access disabled',
                    'status': get_access_status()
                })
            else:
                success, expiry_time = enable_mobile_access(5)
                return jsonify({
                    'success': success,
                    'enabled': True,
                    'message': 'Mobile access enabled for 5 minutes',
                    'expiry_time': expiry_time.isoformat() if success else None,
                    'status': get_access_status()
                })
                
    except Exception as e:
        print(f"[ERROR] Toggle mobile access failed: {e}")
        return jsonify({
            'success': False,
            'message': 'Server error'
        }), 500

@app.route('/api/mobile_access/status', methods=['GET'])
@login_required
def mobile_access_status():
    """Get current mobile access status"""
    try:
        status = get_access_status()
        return jsonify({
            'success': True,
            'status': status
        })
    except Exception as e:
        print(f"[ERROR] Get mobile access status failed: {e}")
        return jsonify({
            'success': False,
            'message': 'Server error'
        }), 500


@app.route('/debug/db_status')
def debug_db_status():
    """Debug endpoint to check database status"""
    if 'username' not in session:
        return redirect(url_for('login'))
    try:
        data = load_attendance_records()
        records = data.get('records', {})

        total_records = 0
        dates = set()
        for key, rec in records.items():
            present = len(rec.get('present', []))
            absent = len(rec.get('absent', []))
            total_records += present + absent
            date = key.split('_', 1)[0]
            dates.add(date)

        today = datetime.now().strftime("%Y-%m-%d")
        today_stats = {'Present': 0, 'Absent': 0}
        for key, rec in records.items():
            if key.startswith(today + "_"):
                today_stats['Present'] = len(rec.get('present', []))
                today_stats['Absent'] = len(rec.get('absent', []))

        return jsonify({
            'tables': ['attendance_records.json', 'student_data.json', 'teachers.json'],
            'total_records': total_records,
            'distinct_dates': len(dates),
            'today_stats': today_stats,
            'lecture': session.get('lecture', '')
        })
    except Exception as e:
        print(f"[ERROR] debug_db_status failed: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/__debug/lecture-data')
@login_required
def debug_lecture_data():
    """Return collect_lecture_data() for the current session (helps debug why /lecture_summary redirects)."""
    try:
        data = collect_lecture_data()
        return jsonify({'lecture_data': data})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/curriculum/state')
def api_curriculum_state():
    try:
        state = get_curriculum_state()
        return jsonify({'success': True, 'state': state})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/curriculum/toggle', methods=['POST'])
@login_required
def api_curriculum_toggle():
    try:
        # toggle_curriculum may return either a side string or a dict depending on
        # which helper implementation is present; always call get_curriculum_state()
        # after toggling to return a consistent structure to the client.
        try:
            new_side = toggle_curriculum()
        except TypeError:
            # defensive: if toggle_curriculum requires arguments, call without
            new_side = toggle_curriculum()

        state = get_curriculum_state()
        return jsonify({'success': True, 'new_side': new_side, 'state': state})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500



# Add these constants near the top with other constants
STUDENT_DATA_JSON = "student_data.json"


# Use load_student_data/save_student_data imported from attendance_system


def get_students_by_subject_and_year(subject, year):
    """Get students filtered by subject and year"""
    student_data = load_student_data()
    filtered_students = []

    for student_id, student_info in student_data.items():
        if (subject in student_info.get('subjects', []) and
                student_info.get('year') == year):
            filtered_students.append({
                'student_id': student_id,
                'name': student_info.get('name', ''),
                'year': student_info.get('year'),
                'image_path': student_info.get('image_path'),
                'batch': student_info.get('batch')
            })

    return filtered_students


def add_student_to_json(student_data):
    """Add a new student to student_data.json"""
    try:
        students = load_student_data()
        student_id = student_data['student_id']

        # Check if student already exists
        if student_id in students:
            return False, "Student ID already exists"

        # Add new student
        students[student_id] = {
            'student_id': student_id,
            'name': student_data['name'],
            'year': student_data['year'],
            'subjects': student_data['subjects'],
            'image_path': student_data.get('image_path', f"student_images/images/photo1762364426.jpg"),
            'major': student_data.get('major', ''),
            'starting_year': student_data.get('starting_year', datetime.now().year)
        }

        # Save back to file, preserving existing on-disk format when possible
        save_student_data(students)

        return True, "Student added successfully"

    except Exception as e:
        return False, f"Error adding student: {str(e)}"


def update_student_in_json(student_id, updated_data):
    """Update existing student data"""
    try:
        students = load_student_data()

        if student_id not in students:
            return False, "Student not found"

        # Update student data
        students[student_id].update(updated_data)

        # Save back to file, preserving existing on-disk format when possible
        save_student_data(students)

        return True, "Student updated successfully"

    except Exception as e:
        return False, f"Error updating student: {str(e)}"


def delete_student_from_json(student_id):
    """Delete student from JSON"""
    try:
        students = load_student_data()

        if student_id not in students:
            return False, "Student not found"

        # Remove student
        del students[student_id]

        # Save back to file, preserving existing on-disk format when possible
        save_student_data(students)

        return True, "Student deleted successfully"

    except Exception as e:
        return False, f"Error deleting student: {str(e)}"


@app.route('/manage_students')
@login_required
def manage_students():
    """Student management interface"""
    try:
        # Prefer filtering by the subject currently selected by the teacher
        selected_subject = session.get('selected_subject') or session.get('lecture')
        teacher_subjects = session.get('teacher_subjects', [])

        # Get students filtered by selected_subject (if present) otherwise fall back to teacher_subjects intersection
        all_students = []
        student_data = load_student_data()

        for student_id, student_info in student_data.items():
            student_subjects = student_info.get('subjects', [])

            if selected_subject:
                # Only include students enrolled in the currently selected subject
                if selected_subject in student_subjects:
                    all_students.append({
                        'student_id': student_id,
                        'name': student_info.get('name', ''),
                        'year': student_info.get('year'),
                        'subjects': student_subjects,
                        'image_path': student_info.get('image_path'),
                        'major': student_info.get('major', ''),
                        'common_subjects': [selected_subject]
                    })
            else:
                # Fallback: include students who share any subject with teacher_subjects
                common_subjects = list(set(teacher_subjects) & set(student_subjects))
                if common_subjects:
                    all_students.append({
                        'student_id': student_id,
                        'name': student_info.get('name', ''),
                        'year': student_info.get('year'),
                        'subjects': student_subjects,
                        'image_path': student_info.get('image_path'),
                        'major': student_info.get('major', ''),
                        'common_subjects': common_subjects
                    })

        # Get teacher photo
        teacher_photo_url = None
        if session.get('teacher_image'):
            teacher_photo_url = url_for('static', filename=f"teacher_images/{session['teacher_image']}")

        return render_template('manage_students.html',
                               teacher_name=session.get('teacher_name', ''),
                               teacher_photo=teacher_photo_url,
                               teacher_subjects=teacher_subjects,
                               students=all_students,
                               current_year=datetime.now().year)

    except Exception as e:
        print(f"[ERROR] Failed to load student management: {e}")
        flash('Failed to load student management', 'error')
        return redirect(url_for('dashboard'))


@app.route('/api/students/<subject>/<int:year>')
@login_required
def get_students_by_subject_year(subject, year):
    """API endpoint to get students by subject and year"""
    try:
        students = get_students_by_subject_and_year(subject, year)
        return jsonify({'success': True, 'students': students})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/students/add', methods=['POST'])
@login_required
def api_add_student():
    """API endpoint to add a new student"""
    try:
        data = request.get_json()

        required_fields = ['student_id', 'name', 'year', 'subjects']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'success': False, 'message': f'Missing required field: {field}'}), 400

        success, message = add_student_to_json(data)

        if success:
            return jsonify({'success': True, 'message': message})
        else:
            return jsonify({'success': False, 'message': message}), 400

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/students/update/<student_id>', methods=['POST'])
@login_required
def api_update_student(student_id):
    """API endpoint to update student data"""
    try:
        data = request.get_json()

        success, message = update_student_in_json(student_id, data)

        if success:
            return jsonify({'success': True, 'message': message})
        else:
            return jsonify({'success': False, 'message': message}), 400

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/students/delete/<student_id>', methods=['DELETE'])
@login_required
def api_delete_student(student_id):
    """API endpoint to delete student"""
    try:
        success, message = delete_student_from_json(student_id)

        if success:
            return jsonify({'success': True, 'message': message})
        else:
            return jsonify({'success': False, 'message': message}), 400

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


# Add these routes to your app.py file (after the /mark_attendance route)

# ---------- ATTENDANCE RECORDS ----------
@app.route('/attendance_records')
def attendance_records():
    """View all attendance records"""
    if 'username' not in session:
        return redirect(url_for('login'))

    try:
        # Get attendance records for current teacher's lecture
        lecture = session.get('lecture', '')
        today = datetime.now().strftime("%Y-%m-%d")
        
        # Load attendance records from JSON
        data = load_attendance_records()
        records = []
        
        # Process records for the current lecture
        for date_lecture, day_data in data.get('records', {}).items():
            date, rec_lecture = date_lecture.split('_', 1)
            if rec_lecture == lecture:
                # Add present students
                for student_id in day_data.get('present', []):
                    student_info = load_student_data().get(student_id, {})
                    records.append((
                        student_id,
                        student_info.get('name', 'Unknown'),
                        date,
                        day_data.get('time', '00:00:00'),
                        'Present',
                        lecture
                    ))
                # Add absent students
                for student_id in day_data.get('absent', []):
                    student_info = load_student_data().get(student_id, {})
                    records.append((
                        student_id,
                        student_info.get('name', 'Unknown'),
                        date,
                        day_data.get('time', '00:00:00'),
                        'Absent',
                        lecture
                    ))

        # Sort records by date and time
        records.sort(key=lambda x: (x[2], x[3]), reverse=True)

        # Count present and absent
        present_count = sum(1 for record in records if record[4] == 'Present')
        absent_count = sum(1 for record in records if record[4] == 'Absent')

        # Get today's statistics
        today_key = f"{today}_{lecture}"
        today_data = data.get('records', {}).get(today_key, {})
        present_today = len(today_data.get('present', []))
        absent_today = len(today_data.get('absent', []))

        # Get teacher photo
        teacher_photo_url = None
        try:
            teachers = load_teachers()
            t = teachers.get(session['username'])
            if t and t.get('photo'):
                teacher_photo_url = url_for('static', filename=f"teacher_images/{t['photo']}")
        except Exception as e:
            print(f"[WARN] Could not load teacher photo: {e}")
            teacher_photo_url = None

        return render_template('attendance_records.html',
                               teacher_name=session.get('teacher_name', ''),
                               teacher_photo=teacher_photo_url,
                               records=records,
                               lecture=lecture,
                               present_count=present_count,
                               absent_count=absent_count,
                               present_today=present_today,
                               absent_today=absent_today)

    except Exception as e:
        print(f"[ERROR] Failed to load attendance records: {e}")
        import traceback
        traceback.print_exc()
        flash('Failed to load attendance records', 'error')
        return redirect(url_for('dashboard'))

@app.route('/finalize_attendance', methods=['POST'])
def finalize_attendance():
    """Finalize attendance for a specific date"""
    if 'username' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'})

    try:
        data = request.get_json()
        date = data.get('date')
        lecture = session.get('lecture', '')

        if not date:
            return jsonify({'success': False, 'message': 'Date required'})

        # Load attendance records
        attendance_data = load_attendance_records()
        date_key = f"{date}_{lecture}"
        
        if date_key in attendance_data.get('records', {}):
            # Save the finalized records
            save_attendance_records(attendance_data)
            return jsonify({
                'success': True, 
                'message': 'Attendance finalized',
                'present_count': len(attendance_data['records'][date_key].get('present', [])),
                'absent_count': len(attendance_data['records'][date_key].get('absent', []))
            })
        else:
            return jsonify({'success': False, 'message': 'No records found for this date'})

    except Exception as e:
        print(f"[ERROR] Finalize attendance failed: {e}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/lecture_summary')
@login_required
def lecture_summary():
    lecture_data = collect_lecture_data()
    if not lecture_data:
        flash('No subject selected yet.', 'error')
        return redirect(url_for('select_subject'))

    return render_template('lecture_summary.html', lecture_data=lecture_data)

@app.route('/clear_attendance', methods=['POST'])
def clear_attendance():
    """Clear attendance records for a specific date"""
    if 'username' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    try:
        data = request.get_json()
        date = data.get('date')
        lecture = session.get('lecture', '')

        if not date:
            return jsonify({'success': False, 'message': 'Date required'}), 400

        attendance = load_attendance_records()
        key = f"{date}_{lecture}"
        if key in attendance.get('records', {}):
            removed = attendance['records'].pop(key)
            save_attendance_records(attendance)
            affected = len(removed.get('present', [])) + len(removed.get('absent', []))
            return jsonify({
                'success': True,
                'message': f'Cleared {affected} records for {date}'
            })
        else:
            return jsonify({'success': True, 'message': 'No records to clear for this date'})

    except Exception as e:
        print(f"[ERROR] Clear attendance failed: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500


@app.route('/clear_and_defaulters', methods=['GET', 'POST'])
def clear_and_defaulters():
    """View clear students and defaulters"""
    if 'username' not in session:
        return redirect(url_for('login'))

    try:
        lecture = session.get('lecture', '')
        threshold = int(request.form.get('threshold', 75)) if request.method == 'POST' else 75

        attendance = load_attendance_records()
        student_data = load_student_data()

        dates = set()
        lecture_records = {}
        for key, rec in attendance.get('records', {}).items():
            parts = key.split('_', 1)
            if len(parts) != 2:
                continue
            date_part, rec_lecture = parts
            if rec_lecture == lecture:
                dates.add(date_part)
                lecture_records[date_part] = rec

        total_classes = len(dates)

        summary = {}
        for date_key, rec in lecture_records.items():
            for sid in rec.get('present', []):
                s = summary.setdefault(sid, {
                    'student_id': sid, 
                    'student_name': student_data.get(sid, {}).get('name', sid), 
                    'present_count': 0
                })
                s['present_count'] += 1
            for sid in rec.get('absent', []):
                summary.setdefault(sid, {
                    'student_id': sid, 
                    'student_name': student_data.get(sid, {}).get('name', sid), 
                    'present_count': 0
                })

        students = []
        for sid, s in summary.items():
            present_count = s['present_count']
            # FIX: Use total_classes (conducted) instead of individual student's total
            percentage = round((present_count * 100.0 / total_classes), 2) if total_classes > 0 else 0.0
            students.append((s['student_id'], s['student_name'], present_count, percentage))

        students.sort(key=lambda x: x[3], reverse=True)  # Sort by percentage (index 3)
        total_students = len(students)
        clear_students = [s for s in students if s[3] >= threshold]  # Index 3 is percentage
        defaulters = [s for s in students if s[3] < threshold]

        # Get teacher photo
        teacher_photo_url = None
        try:
            teachers = load_teachers()
            t = teachers.get(session['username'])
            if t and t.get('photo'):
                teacher_photo_url = url_for('static', filename=f"teacher_images/{t['photo']}")
        except Exception:
            teacher_photo_url = None

        return render_template('clear_and_defaulters.html',
                               teacher_name=session.get('teacher_name', ''),
                               teacher_photo=teacher_photo_url,
                               lecture=lecture,
                               total_classes=total_classes,
                               threshold=threshold,
                               total_students=total_students,
                               clear_students=clear_students,
                               defaulters=defaulters)
    except Exception as e:
        return f"Error: {str(e)}", 500


@app.route('/get_defaulters', methods=['GET'])
def get_defaulters():
    """Get list of students with low attendance"""
    if 'username' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    try:
        lecture = session.get('lecture', '')
        threshold = int(request.args.get('threshold', 75))

        attendance = load_attendance_records()
        student_data = load_student_data()

        summary = {}
        for key, rec in attendance.get('records', {}).items():
            parts = key.split('_', 1)
            if len(parts) != 2:
                continue
            _, rec_lecture = parts
            if rec_lecture != lecture:
                continue
            for sid in rec.get('present', []):
                s = summary.setdefault(sid, {'student_id': sid, 'student_name': student_data.get(sid, {}).get('name',''), 'total_classes':0, 'present_count':0})
                s['total_classes'] += 1
                s['present_count'] += 1
            for sid in rec.get('absent', []):
                s = summary.setdefault(sid, {'student_id': sid, 'student_name': student_data.get(sid, {}).get('name',''), 'total_classes':0, 'present_count':0})
                s['total_classes'] += 1

        defaulters = []
        for sid, s in summary.items():
            if s['total_classes'] == 0:
                continue
            percentage = round((s['present_count'] * 100.0 / s['total_classes']), 2)
            if percentage < threshold:
                defaulters.append({
                    'student_id': sid,
                    'student_name': s['student_name'],
                    'total_classes': s['total_classes'],
                    'present_count': s['present_count'],
                    'percentage': percentage
                })

        defaulters.sort(key=lambda x: x['percentage'])
        return jsonify({
            'success': True,
            'defaulters': defaulters,
            'threshold': threshold
        })

    except Exception as e:
        print(f"[ERROR] Get defaulters failed: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500

@app.route('/download_clear_defaulters/<filetype>/<int:threshold>')
def download_clear_defaulters(filetype, threshold):
    """Download clear students and defaulters report"""
    if 'username' not in session:
        return redirect(url_for('login'))

    try:
        lecture = session.get('lecture', '')
        teacher_name = session.get('teacher_name', '')

        attendance = load_attendance_records()
        student_data = load_student_data()

        summary = {}
        for key, rec in attendance.get('records', {}).items():
            parts = key.split('_', 1)
            if len(parts) != 2:
                continue
            _, rec_lecture = parts
            if rec_lecture != lecture:
                continue
            for sid in rec.get('present', []):
                s = summary.setdefault(sid, {'student_id': sid, 'student_name': student_data.get(sid, {}).get('name',''), 'total_classes':0, 'present_count':0})
                s['total_classes'] += 1
                s['present_count'] += 1
            for sid in rec.get('absent', []):
                s = summary.setdefault(sid, {'student_id': sid, 'student_name': student_data.get(sid, {}).get('name',''), 'total_classes':0, 'present_count':0})
                s['total_classes'] += 1

        students = []
        for sid, s in summary.items():
            percentage = round((s['present_count'] * 100.0 / s['total_classes']), 2) if s['total_classes'] > 0 else 0.0
            students.append((s['student_id'], s['student_name'], s['total_classes'], s['present_count'], percentage))

        students.sort(key=lambda x: x[4], reverse=True)
        clear_students = [s for s in students if s[4] >= threshold]
        defaulters = [s for s in students if s[4] < threshold]

        if filetype == 'excel':
            # Create Excel workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Attendance Summary"

            # Title
            ws.merge_cells('A1:D1')
            ws['A1'] = f"Attendance Summary - {lecture}"
            ws['A1'].font = Font(bold=True, size=14)
            ws['A1'].alignment = Alignment(horizontal='center')

            ws.merge_cells('A2:D2')
            ws['A2'] = f"Teacher: {teacher_name} | Threshold: {threshold}%"
            ws['A2'].alignment = Alignment(horizontal='center')

            # Clear Students
            ws['A4'] = "Clear Students (≥ {}%)".format(threshold)
            ws['A4'].font = Font(bold=True)

            headers = ['Student ID', 'Student Name', 'Classes Attended', 'Attendance %']
            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=5, column=col)
                cell.value = header
                cell.font = Font(bold=True)

            for row_idx, student in enumerate(clear_students, start=6):
                for col_idx, value in enumerate(student[:4], start=1):
                    ws.cell(row=row_idx, column=col_idx).value = value

            # Defaulters
            start_row = len(clear_students) + 8
            ws.cell(row=start_row, column=1).value = "Defaulters (< {}%)".format(threshold)
            ws.cell(row=start_row, column=1).font = Font(bold=True)

            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=start_row+1, column=col)
                cell.value = header
                cell.font = Font(bold=True)

            for row_idx, student in enumerate(defaulters, start=start_row+2):
                for col_idx, value in enumerate(student[:4], start=1):
                    ws.cell(row=row_idx, column=col_idx).value = value

            # Save to BytesIO
            output = BytesIO()
            wb.save(output)
            output.seek(0)

            filename = f"attendance_summary_{lecture}_{datetime.now().strftime('%Y%m%d')}.xlsx"
            return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                           as_attachment=True, download_name=filename)

        else:
            flash('Invalid file type', 'error')
            return redirect(url_for('clear_and_defaulters'))

    except Exception as e:
        print(f"[ERROR] Download failed: {e}")
        flash('Download failed', 'error')
        return redirect(url_for('clear_and_defaulters'))

@app.route('/export_attendance/<format>')
def export_attendance(format):
    """Export attendance records to Excel or Word"""
    if 'username' not in session:
        return redirect(url_for('login'))

    try:
        lecture = session.get('lecture', '')
        teacher_name = session.get('teacher_name', '')

        # Read from JSON-based attendance records
        attendance = load_attendance_records()
        student_data = load_student_data()

        records = []
        for key, rec in attendance.get('records', {}).items():
            try:
                date_part, rec_lecture = key.split('_', 1)
            except Exception:
                continue
            if rec_lecture != lecture:
                continue
            time_str = rec.get('time', '')
            for sid in rec.get('present', []):
                records.append((sid, student_data.get(sid, {}).get('name', ''), date_part, time_str, 'Present'))
            for sid in rec.get('absent', []):
                records.append((sid, student_data.get(sid, {}).get('name', ''), date_part, time_str, 'Absent'))

        # Sort records by date desc then student_id asc
        records.sort(key=lambda x: (x[2], x[0]), reverse=True)

        if format == 'excel':
            # Create Excel workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Attendance Records"

            # Styling
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=12)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # Title
            ws.merge_cells('A1:E1')
            title_cell = ws['A1']
            title_cell.value = f"Attendance Records - {lecture}"
            title_cell.font = Font(bold=True, size=14)
            title_cell.alignment = Alignment(horizontal='center')

            ws.merge_cells('A2:E2')
            subtitle_cell = ws['A2']
            subtitle_cell.value = f"Teacher: {teacher_name}"
            subtitle_cell.alignment = Alignment(horizontal='center')

            # Headers
            headers = ['Student ID', 'Student Name', 'Date', 'Time', 'Status']
            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=4, column=col)
                cell.value = header
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center')
                cell.border = border

            # Data
            for row_idx, record in enumerate(records, start=5):
                for col_idx, value in enumerate(record, start=1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell.value = value
                    cell.border = border
                    cell.alignment = Alignment(horizontal='center')

                    # Color code status
                    if col_idx == 5:  # Status column
                        if value == 'Present':
                            cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
                        else:
                            cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")

            # Adjust column widths
            ws.column_dimensions['A'].width = 15
            ws.column_dimensions['B'].width = 25
            ws.column_dimensions['C'].width = 15
            ws.column_dimensions['D'].width = 15
            ws.column_dimensions['E'].width = 12

            # Save to BytesIO
            output = BytesIO()
            wb.save(output)
            output.seek(0)

            filename = f"attendance_{lecture}_{datetime.now().strftime('%Y%m%d')}.xlsx"
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=filename
            )

        elif format == 'word':
            # Create Word document
            doc = Document()

            # Title
            doc.add_heading(f'Attendance Records - {lecture}', 0)
            doc.add_paragraph(f'Teacher: {teacher_name}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()

            # Table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            # Headers
            headers = ['Student ID', 'Student Name', 'Date', 'Time', 'Status']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header

            # Data
            for record in records:
                row_cells = table.add_row().cells
                for i, value in enumerate(record):
                    row_cells[i].text = str(value)

            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)

            filename = f"attendance_{lecture}_{datetime.now().strftime('%Y%m%d')}.docx"
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=filename
            )

        else:
            flash('Invalid export format', 'error')
            return redirect(url_for('attendance_records'))

    except Exception as e:
        print(f"[ERROR] Export failed: {e}")
        flash('Export failed', 'error')
        return redirect(url_for('attendance_records'))


@app.route('/sync_database', methods=['POST'])
def sync_database():
    """Sync attendance data between JSON and database"""
    if 'username' not in session:
        return jsonify({'success': False, 'message': 'Not logged in'}), 401

    try:
        # This is a placeholder - implement based on your needs
        # You might sync data from attendance_system.json to database

        return jsonify({
            'success': True,
            'message': 'Database synchronized successfully'
        })

    except Exception as e:
        print(f"[ERROR] Sync failed: {e}")
        return jsonify({'success': False, 'message': 'Server error'}), 500



# -----------------------------
# SECTION: Video control & attendance process
# (start/stop endpoints, spawn attendance thread, video feed)
# -----------------------------

# ---------- VIDEO CONTROL ----------
@app.route('/start', methods=['POST'])
def start_script():
    global process_thread
    if not app.shared_data.get('running', False):
        app.shared_data['running'] = True

        def run_attendance_system():
            try:
                attendance_main(app.shared_data)
            except Exception as e:
                print("[ERROR] attendance_main crashed:", e)
            finally:
                app.shared_data['running'] = False

        process_thread = threading.Thread(target=run_attendance_system, daemon=True)
        process_thread.start()
        print("[INFO] Attendance system thread started.")
    return 'Started'


@app.route('/stop', methods=['POST'])
def stop_script():
    if app.shared_data.get('running', False):
        app.shared_data['running'] = False
        print("[INFO] Stopping attendance system...")

        # Auto-mark absent students with delay to ensure all present records are saved
        try:
            lecture = session.get('lecture', '') or session.get('selected_subject', '')
            if lecture:
                # Wait for any pending face recognition to complete
                import time
                time.sleep(5)  # Increased to 5 seconds to ensure all processing completes
                mark_absent_for_unmarked_students(lecture, session.get('username', 'system'))
                print(f"[INFO] Auto-marked absent students for lecture: {lecture}")
        except Exception as e:
            print(f"[ERROR] Failed to auto-mark absent: {e}")

    return 'Stopped'


@app.route('/debug/student/<student_id>')
def debug_student(student_id):
    """Debug endpoint to check student data"""
    try:
        students = load_student_data()
        student = students.get(student_id)

        attendance_data = load_attendance_records()
        recent = []
        for key, rec in attendance_data.get('records', {}).items():
            date_part, lecture = key.split('_', 1)
            if student_id in rec.get('present', []) or student_id in rec.get('absent', []):
                status = 'Present' if student_id in rec.get('present', []) else 'Absent'
                recent.append((student_id, student.get('name') if student else 'Unknown', date_part, rec.get('time', ''), status))

        recent.sort(key=lambda x: x[2], reverse=True)
        recent = recent[:5]

        return jsonify({
            'student': student,
            'recent_attendance': recent
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/debug/tables')
def debug_tables():
    """Debug endpoint to check all tables"""
    try:
        attendance = load_attendance_records()
        students = load_student_data()
        teachers = load_teachers()

        table_data = {
            'attendance_records.json': {
                'count': sum(len(rec.get('present', [])) + len(rec.get('absent', [])) for rec in attendance.get('records', {}).values()),
                'sample': list(attendance.get('records', {}).items())[:3]
            },
            'student_data.json': {
                'count': len(students),
                'sample': list(students.items())[:3]
            },
            'teachers.json': {
                'count': len(teachers),
                'sample': list(teachers.items())[:3]
            }
        }
        return jsonify(table_data)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/debug/attendance_status')
def debug_attendance_status():
    """Debug endpoint to check current attendance status"""
    if 'username' not in session and 'teacher_id' not in session:
        return jsonify({'error': 'Not logged in'}), 401

    try:
        lecture = session.get('lecture', '') or session.get('selected_subject', '')
        today = datetime.now().strftime("%Y-%m-%d")

        attendance = load_attendance_records()
        student_data = load_student_data()

        key = f"{today}_{lecture}"
        rec = attendance.get('records', {}).get(key, {})
        present = rec.get('present', [])
        absent = rec.get('absent', [])

        records = []
        for sid in present:
            records.append((sid, student_data.get(sid, {}).get('name', ''), 'Present', rec.get('time', '')))
        for sid in absent:
            records.append((sid, student_data.get(sid, {}).get('name', ''), 'Absent', rec.get('time', '')))

        present_count = len(present)
        absent_count = len(absent)

        all_students = [{'student_id': sid, 'student_name': s.get('name','')} for sid, s in student_data.items()]

        return jsonify({
            'lecture': lecture,
            'date': today,
            'total_students': len(all_students),
            'present_count': present_count,
            'absent_count': absent_count,
            'records': [
                {'student_id': r[0], 'student_name': r[1], 'status': r[2], 'time': r[3]} for r in records
            ],
            'all_students': all_students
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


def mark_absent_for_unmarked_students(lecture, marked_by="system"):
    """Mark all students who didn't attend TODAY as absent - USING STUDENT_DATA.JSON"""
    today = datetime.now().strftime("%Y-%m-%d")
    current_time = datetime.now().strftime("%H:%M:%S")
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    try:
        # Get students enrolled in this lecture from student_data.json
        student_data = load_student_data()
        enrolled_students = []

        for student_id, student_info in student_data.items():
            if lecture in student_info.get('subjects', []):
                enrolled_students.append({
                    'student_id': student_id,
                    'name': student_info.get('name', 'Unknown')
                })

        if not enrolled_students:
            print(f"[WARN] No students found enrolled in {lecture}")
            return

        # Load attendance JSON and mark absent where necessary
        attendance = load_attendance_records()
        key = f"{today}_{lecture}"
        if key not in attendance.get('records', {}):
            attendance.setdefault('records', {})[key] = {'present': [], 'absent': [], 'time': current_time}

        today_rec = attendance['records'][key]
        present_students = set(today_rec.get('present', []))
        absent_students = set(today_rec.get('absent', []))

        absent_count = 0
        for student in enrolled_students:
            sid = student['student_id']
            if sid not in present_students and sid not in absent_students:
                today_rec.setdefault('absent', []).append(sid)
                absent_count += 1

        save_attendance_records(attendance)
        print(f"[SUCCESS] Marked {absent_count} students as absent for {lecture} on {today}")
    except Exception as e:
        print(f"[ERROR] Failed to mark absent students for {lecture}: {e}")



def generate_frames():
    while True:
        if not app.shared_data.get('running', False):
            time.sleep(0.1)
            continue

        frame_bytes = app.shared_data.get('frame_jpeg', None)
        if frame_bytes:
            try:
                yield (b'--frame\r\n'
                       b'Content-Type: image/jpeg\r\n\r\n' + frame_bytes + b'\r\n')
            except Exception as e:
                print("[ERROR] Failed to yield frame:", e)
                time.sleep(0.05)
        else:
            time.sleep(0.05)


@app.route('/video_feed')
def video_feed():
    def generate():
        last_frame_time = 0
        
        while True:
            if 'frame_jpeg' in shared_data and 'frame_updated' in shared_data:
                if shared_data['frame_updated'] > last_frame_time:
                    last_frame_time = shared_data['frame_updated']
                    
                    yield (b'--frame\r\n'
                           b'Content-Type: image/jpeg\r\n\r\n' + 
                           shared_data['frame_jpeg'] + b'\r\n')
            
            time.sleep(0.033)  # Check 30 times/sec - FASTER CHECKS
    
    return Response(generate(),
                    mimetype='multipart/x-mixed-replace; boundary=frame')



@app.route('/student_info', methods=['GET'])
@login_required
def get_student_info():
    """Read current_student.json to display on frontend"""
    try:
        with open("current_student.json", "r", encoding='utf-8') as f:
            data = json.load(f)
        return jsonify(data)
    except FileNotFoundError:
        return jsonify({"error": "No student recognized yet."})
    except Exception as e:
        print(f"[ERROR] Failed to read current_student.json: {e}")
        return jsonify({"error": "Failed to load student info"})

def run_flask_app(shared_data):
    app.shared_data = shared_data
    app.run(debug=True, use_reloader=False, host="0.0.0.0", port=5000)

@app.route('/')
def home():
    client_ip = request.remote_addr
    
    # Check if it's localhost (teacher's device)
    if is_localhost(client_ip):
        # Teacher accessing from their laptop
        if 'username' in session or 'teacher_id' in session:
            return redirect(url_for('dashboard'))
        else:
            return redirect(url_for('login'))
    
    # Check if it's a network device (student's phone/tablet)
    elif is_network_device(client_ip):
        # Student accessing from their device
        # Get current teacher and subject info
        try:
            with open('current_teacher.json', 'r') as f:
                teacher_data = json.load(f)
        except:
            teacher_data = {
                'name': 'Teacher',
                'lecture': 'Unknown Subject'
            }
        
        # Serve mobile attendance page directly
        return render_template('mobile_attendance.html', 
                              teacher_name=teacher_data.get('name', 'Teacher'),
                              subject=teacher_data.get('lecture', 'Unknown Subject'),
                              server_ip=SERVER_LAN_IP)
    
    # Fallback (should rarely happen)
    else:
        return redirect(url_for('login'))

register_mobile_routes(app)

# -----------------------------
# SECTION: Main entrypoint
# (process startup, multiprocessing manager and Flask run)
# -----------------------------

if __name__ == '__main__':
    multiprocessing.freeze_support()
    manager = multiprocessing.Manager()
    shared_data = manager.dict()
    shared_data['running'] = False
    shared_data['frame_jpeg'] = None

    # Initialize database
    init_database()

    run_flask_app(shared_data)